import os
import glob
import re
import markdown
from bs4 import BeautifulSoup
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_UNDERLINE

def preprocess_markdown(md_content):
    lines = md_content.split('\n')
    new_lines = []
    pattern_4stars = re.compile(r'^\*\*\*\*(.+)\*\*\*\*$')
    pattern_2stars = re.compile(r'^\*\*(.+)\*\*$')
    pattern_title_colon = re.compile(r'^\*\*(.+)\*\*\s*:\s*(.*)$')

    for line in lines:
        stripped = line.strip()
        match_4 = pattern_4stars.match(stripped)
        if match_4:
            new_lines.append(f"# {match_4.group(1).strip()}")
            continue

        match_2 = pattern_2stars.match(stripped)
        if match_2:
            new_lines.append(f"# {match_2.group(1).strip()}")
            continue

        match_title_colon = pattern_title_colon.match(stripped)
        if match_title_colon:
            new_lines.append(f"**{match_title_colon.group(1).strip()}: {match_title_colon.group(2).strip()}**")
            continue

        new_lines.append(line)
    
    return "\n".join(new_lines)

def create_first_page(doc, first_image_1, first_image_2):
    doc.add_picture(first_image_1, width=Inches(8.5))
    heading = doc.add_paragraph()
    heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = heading.add_run('Proposal')
    run.font.size = Pt(36)
    run.font.name = 'Arial'
    run.font.bold = True
    
    for text in ['To', '| logo |', 'for', '| Project Title |']:
        para = doc.add_paragraph()
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = para.add_run(text)
        run.font.size = Pt(14)
    
    doc.add_picture(first_image_2, width=Inches(8.5))

def process_markdown_files(doc, md_files):
    for md_file in md_files:
        try:
            with open(md_file, 'r', encoding='utf-8') as f:
                original_md_content = f.read()
            
            preprocessed_content = preprocess_markdown(original_md_content)
            
        except UnicodeDecodeError:
            with open(md_file, 'r', encoding='utf-8-sig') as f:
                original_md_content = f.read()
            
            preprocessed_content = preprocess_markdown(original_md_content)
        
        html_content = markdown.markdown(preprocessed_content, extensions=["tables"])
        soup = BeautifulSoup(html_content, "html.parser")
        
        for element in soup.children:
            if element.name in ["h1", "h2", "h3"]:
                p = doc.add_paragraph()
                run = p.add_run(element.get_text())
                run.font.bold = True
                run.font.size = Pt(14) if element.name == "h1" else Pt(12)
            elif element.name == "p":
                p = doc.add_paragraph(element.get_text())
            elif element.name == "ul" or element.name == "ol":
                for li in element.find_all("li"):
                    doc.add_paragraph(li.get_text(), style="List Bullet")
            elif element.name == "table":
                rows = element.find_all("tr")
                if rows:
                    table = doc.add_table(rows=len(rows), cols=len(rows[0].find_all(["th", "td"])))
                    for row_idx, row in enumerate(rows):
                        for col_idx, cell in enumerate(row.find_all(["th", "td"])):
                            table.cell(row_idx, col_idx).text = cell.get_text()
        
        doc.add_paragraph("\n")

def create_docx(md_files, output_docx, logo_path, first_image_1, first_image_2):
    doc = Document()
    create_first_page(doc, first_image_1, first_image_2)
    process_markdown_files(doc, md_files)
    
    # --- Add Footer ---
    footer = doc.sections[0].footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = footer.add_run("Copyright © 2025, Nitor Infotech Pvt. Ltd., All Rights Reserved. Confidential")
    run.font.size = Pt(10)
    
    doc.save(output_docx)

if __name__ == "__main__":
    md_files = glob.glob(r"server/app/pipelines/proposal_generator/proposals/*.md")
    logo_path = r"server/app/pipelines/proposal_generator/nitor_logo.png"
    first_image_1 = r"server/app/pipelines/proposal_generator/first_image_1.png"
    first_image_2 = r"server/app/pipelines/proposal_generator/first_image_2.png"
    create_docx(md_files, "Nitor_proposal1.docx", logo_path, first_image_1, first_image_2)
